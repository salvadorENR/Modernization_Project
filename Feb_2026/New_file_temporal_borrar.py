import os
import sys
import datetime
import io
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import seaborn as sns
import scipy.stats as stats

# Intentar importar docx para crear el Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("¡ERROR! Falta instalar la librería para Word.")
    print("Abre tu terminal y ejecuta: py -m pip install python-docx")
    sys.exit(1)

# =========================================================================
# 1. CONFIGURACIÓN INICIAL Y DIRECTORIOS
# =========================================================================
script_dir = os.path.dirname(os.path.abspath(__file__))
dir_txt = os.path.join(script_dir, "Analisis_Diagnostico")
dir_pdf = os.path.join(script_dir, "Reporte_Tecnico")

os.makedirs(dir_txt, exist_ok=True)
os.makedirs(dir_pdf, exist_ok=True)

archivo_pdf = os.path.join(dir_pdf, "Reporte_Tecnico_100_Completos.pdf")
archivo_docx = os.path.join(dir_pdf, "Reporte_Tecnico_Editable.docx")
archivo_txt_global = os.path.join(dir_txt, "Reporte_Finalizacion_100.txt")

global_totals = {
    'Matemática': {'Registrados': 0, '100%_Completos': 0},
    'Lengua': {'Registrados': 0, '100%_Completos': 0}
}
df_global_estand = pd.DataFrame()

with open(archivo_txt_global, 'w', encoding='utf-8') as f:
    f.write("=================================================================\n")
    f.write(" REPORTE DE COMPLETITUD Y FILTRADO 100% (ANÁLISIS INDEPENDIENTE)\n")
    f.write("=================================================================\n\n")

def print_txt(texto):
    print(texto, end='')
    with open(archivo_txt_global, 'a', encoding='utf-8') as f:
        f.write(texto)

# =========================================================================
# 2. FUNCIONES DE PROCESAMIENTO
# =========================================================================
def leer_y_limpiar(filepath):
    if not os.path.exists(filepath): return None
    try:
        df = pd.read_csv(filepath, sep='|', encoding='utf-8', dtype=str, on_bad_lines='skip')
    except UnicodeDecodeError:
        df = pd.read_csv(filepath, sep='|', encoding='latin-1', dtype=str, on_bad_lines='skip')
    
    df.columns = [c.replace('"', '').strip() for c in df.columns]
    if 'Documento' not in df.columns: return None
    
    df['Documento'] = df['Documento'].str.strip()
    if 'Centro' in df.columns:
        df['Centro'] = df['Centro'].str.strip().replace('', 'DESCONOCIDO').fillna('DESCONOCIDO')
    else:
        df['Centro'] = 'DESCONOCIDO'
        
    df.dropna(subset=['Documento'], inplace=True)
    df = df[df['Documento'] != '']
    df.drop_duplicates(subset=['Documento'], keep='first', inplace=True)
    return df

def procesar_100_porciento(df, materia, prefijo):
    if df is None or df.empty: return None
        
    col_pt = [c for c in df.columns if 'Puntaje' in c and 'Total' in c]
    if not col_pt: return None
    col_pt = col_pt[0]
    
    item_cols = [c for c in df.columns if c.startswith(prefijo) and c[len(prefijo):].isdigit()]
    tot_items = len(item_cols)
    if tot_items == 0: return None
    
    valid_mask = df[item_cols].isin(['0', '1']).sum(axis=1) == tot_items
    df_100 = df[valid_mask].copy()
    
    df_100['Puntaje_Total'] = df_100[col_pt].str.replace(',', '.').astype(float)
    
    tot_inicial = len(df)
    tot_final = len(df_100)
    pct = (tot_final / tot_inicial) * 100 if tot_inicial > 0 else 0
    
    global_totals[materia]['Registrados'] += tot_inicial
    global_totals[materia]['100%_Completos'] += tot_final
    
    print_txt(f" [{materia.upper()}]\n")
    print_txt(f"   - Total de ítems en la prueba : {tot_items}\n")
    print_txt(f"   - Estudiantes Registrados     : {tot_inicial:,}\n")
    print_txt(f"   - Completaron el 100%         : {tot_final:,} ({pct:.1f}%)\n\n")
    
    participacion = {'Registrados': tot_inicial, 'Completos': tot_final, 'Pct': pct}
    return {'df': df_100, 'items': item_cols, 'participacion': participacion}

def extraer_metricas(datos_obj, materia, grado_num):
    global df_global_estand
    if datos_obj is None or datos_obj['df'].empty: return None
    
    df_100 = datos_obj['df']
    item_cols = datos_obj['items']
    participacion = datos_obj['participacion']
    
    media = df_100['Puntaje_Total'].mean()
    sd = df_100['Puntaje_Total'].std(ddof=1)
    if pd.isna(sd) or sd == 0: sd = 1
    
    # Manejo seguro de Simetría y Curtosis
    try:
        simetria = float(np.round(stats.skew(df_100['Puntaje_Total'].dropna()), 2))
        if np.isnan(simetria): simetria = 0.0
    except:
        simetria = 0.0
        
    try:
        curtosis = float(np.round(stats.kurtosis(df_100['Puntaje_Total'].dropna()), 2))
        if np.isnan(curtosis): curtosis = 0.0
    except:
        curtosis = 0.0
    
    # Estandarización: Z*1 + 5
    df_100['Z_Score'] = (df_100['Puntaje_Total'] - media) / sd
    df_100['Puntaje_Transformado'] = df_100['Z_Score'] + 5 
    
    df_temp = df_100[['Centro', 'Documento', 'Puntaje_Transformado']].copy()
    df_temp['Materia'] = materia
    df_temp['Grado_Num'] = grado_num
    df_global_estand = pd.concat([df_global_estand, df_temp], ignore_index=True)
    
    desc = pd.DataFrame([{
        'Materia': materia,
        'N (100%)': f"{len(df_100):,}",
        'Media': round(media, 2),
        'Mediana': df_100['Puntaje_Total'].median(),
        'DS': round(sd, 2),
        'Mínimo': df_100['Puntaje_Total'].min(),
        'Máximo': df_100['Puntaje_Total'].max(),
        'Simetría': simetria,
        'Curtosis': curtosis
    }])
    
    # Manejo de Quintiles
    quint_list = []
    if len(df_100) >= 5:
        df_100['Quintil'] = pd.qcut(df_100['Puntaje_Total'].rank(method='first'), 5, labels=[1, 2, 3, 4, 5])
        for q, g in df_100.groupby('Quintil', observed=True):
            quint_list.append({
                'Materia': materia,
                'Quintil': q,
                'Rango': f"{g['Puntaje_Total'].min():.1f} - {g['Puntaje_Total'].max():.1f}",
                'Número Estudiantes': f"{len(g):,}",
                'Porcentaje (%)': f"{(len(g)/len(df_100))*100:.1f}%"
            })
    else:
        quint_list.append({
            'Materia': materia,
            'Quintil': 'Todos',
            'Rango': f"{df_100['Puntaje_Total'].min():.1f} - {df_100['Puntaje_Total'].max():.1f}",
            'Número Estudiantes': f"{len(df_100):,}",
            'Porcentaje (%)': "100.0%"
        })
    quint = pd.DataFrame(quint_list)
    
    esc_z = df_100.groupby('Centro')['Puntaje_Total'].mean().reset_index()
    esc_mean = esc_z['Puntaje_Total'].mean()
    esc_sd = esc_z['Puntaje_Total'].std(ddof=1)
    if pd.isna(esc_sd) or esc_sd == 0: esc_sd = 1
    esc_z['Z_Score'] = (esc_z['Puntaje_Total'] - esc_mean) / esc_sd
    
    dif_series = (df_100[item_cols] == '1').mean() * 100
    dif = pd.DataFrame({'Item': dif_series.index, 'Pct': dif_series.values})
    
    # Retornamos el Z_Score de los estudiantes para estandarizar sus histogramas
    return {'desc': desc, 'quint': quint, 'esc_z': esc_z, 'dif': dif, 'part': participacion, 'df_scores': df_100['Z_Score']}

# =========================================================================
# 3. EXTRAER DATOS MAESTROS
# =========================================================================
print("\nAnalizando bases de datos y procesando filtrado 100%...\n")
grados_data = {}
for grado_num in range(3, 12):
    archivo_mat = os.path.join(script_dir, f"Mat_{grado_num}.txt")
    archivo_lec = os.path.join(script_dir, f"Lec_{grado_num}.txt")
    
    df_mat_raw = leer_y_limpiar(archivo_mat) if os.path.exists(archivo_mat) else None
    df_lec_raw = leer_y_limpiar(archivo_lec) if os.path.exists(archivo_lec) else None
    
    if df_mat_raw is None and df_lec_raw is None: continue
        
    print_txt(f"-> Procesando Grado {grado_num}...\n")
    datos_mat = procesar_100_porciento(df_mat_raw, "Matemática", "MAT")
    datos_lec = procesar_100_porciento(df_lec_raw, "Lengua", "LEC")
    
    met_mat = extraer_metricas(datos_mat, "Matemática", grado_num)
    met_lec = extraer_metricas(datos_lec, "Lengua", grado_num)
    
    if not met_mat and not met_lec: continue
    grados_data[grado_num] = {'mat': met_mat, 'lec': met_lec}

# Textos descriptivos (Enfoque de muestra, sin mencionar "nacional")
txt_boxplot = "El siguiente gráfico presenta la distribución de los puntajes estandarizados (escala con media 5\ny desviación estándar 1, agrupando la mayor parte de los datos entre 2 y 8) para todos los grados\nevaluados. Las cajas muestran el rango intercuartílico, permitiendo observar la variabilidad y la\ntendencia central de la muestra evaluada."

txt_quintiles = "La distribución por quintiles divide a la población estudiantil en cinco grupos de igual tamaño\n(20% cada uno) ordenados según su puntaje. Los rangos muestran el nivel de desempeño alcanzado\nen cada estrato."

txt_hist_alum = "Los histogramas muestran la distribución estandarizada (Z-score) de los puntajes obtenidos por todos los\nestudiantes del grado. Un valor de 0 indica que el alumno se encuentra exactamente en el promedio general."

txt_hist = "Los histogramas muestran la distribución de los promedios de cada Centro Educativo en puntajes Z.\nUn valor de 0 indica que la escuela se encuentra exactamente en el promedio global del grado.\nLas barras ilustran la concentración de las escuelas evaluadas."

# =========================================================================
# 4. CREACIÓN DEL PDF
# =========================================================================
def format_table_pdf(ax, df, title, bold_col_name=None):
    ax.axis('tight')
    ax.axis('off')
    ax.set_title(title, fontweight='bold', pad=12, fontsize=12)
    table = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.7) 
    bold_col_idx = list(df.columns).index(bold_col_name) if bold_col_name in df.columns else -1
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor('darkblue')
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')
        elif col == bold_col_idx:
            cell.get_text().set_fontweight('bold')

print("Generando Documento PDF...\n")
with PdfPages(archivo_pdf) as pdf:
    # --- PORTADA PDF ---
    fig = plt.figure(figsize=(11, 8.5))
    plt.axis('off')
    plt.text(0.5, 0.65, "RESULTADOS DESCRIPTIVOS\nDE LA APLICACIÓN", ha='center', va='center', fontsize=26, fontweight='bold', color='darkblue')
    plt.text(0.5, 0.45, "PRUEBA DE RESULTADOS", ha='center', va='center', fontsize=22)
    plt.text(0.5, 0.35, "FEBRERO 2026", ha='center', va='center', fontsize=14, color='gray', style='italic')
    pdf.savefig(fig)
    plt.close()

    # --- RESUMEN EJECUTIVO GLOBAL PDF ---
    fig = plt.figure(figsize=(11, 8.5))
    fig.suptitle("RESUMEN EJECUTIVO GLOBAL", fontsize=16, fontweight='bold', color='darkblue', y=0.96)
    
    # NUEVAS COORDENADAS PARA EVITAR SUPERPOSICIÓN:
    # La tabla inicia en Y=0.80
    ax_table = fig.add_axes([0.1, 0.80, 0.8, 0.10])
    global_df = pd.DataFrame([
        {'Materia': 'Matemática', 'Total Registrados': f"{global_totals['Matemática']['Registrados']:,}", 'Completaron 100%': f"{global_totals['Matemática']['100%_Completos']:,}", 'Porcentaje': f"{(global_totals['Matemática']['100%_Completos']/max(1, global_totals['Matemática']['Registrados']))*100:.1f}%"},
        {'Materia': 'Lengua', 'Total Registrados': f"{global_totals['Lengua']['Registrados']:,}", 'Completaron 100%': f"{global_totals['Lengua']['100%_Completos']:,}", 'Porcentaje': f"{(global_totals['Lengua']['100%_Completos']/max(1, global_totals['Lengua']['Registrados']))*100:.1f}%"}
    ])
    format_table_pdf(ax_table, global_df, "Tasa de Participación Global (Todos los Grados)")

    # El texto inicia en Y=0.74, totalmente separado de la tabla
    fig.text(0.5, 0.74, txt_boxplot, ha='center', va='top', fontsize=10, style='italic', color='black')

    if not df_global_estand.empty:
        materias_presentes = df_global_estand['Materia'].unique()
        num_mats = len(materias_presentes)
        
        # Los gráficos se ajustaron a 0.24 de altura y bajan hasta 0.36 y 0.05 respectivamente
        # para nunca tocar el texto superior.
        if num_mats == 2:
            axes = [fig.add_axes([0.1, 0.36, 0.8, 0.24]), fig.add_axes([0.1, 0.05, 0.8, 0.24])]
        else:
            axes = [fig.add_axes([0.1, 0.1, 0.8, 0.50])]
            
        for i, mat in enumerate(materias_presentes):
            ax = axes[i]
            df_plot = df_global_estand[df_global_estand['Materia'] == mat]
            sns.boxplot(data=df_plot, x='Grado_Num', y='Puntaje_Transformado', color='lightgray', ax=ax)
            ax.set_title(f"Distribución {mat.upper()} (Z+5)", fontweight='bold')
            ax.set_xlabel("Grado Analizado")
            ax.set_ylabel("Puntaje Estandarizado")
            ax.set_ylim(0, 10)
            ax.grid(axis='y', linestyle='--', alpha=0.7)

    pdf.savefig(fig)
    plt.close()

    # (El resto del código se mantiene igual, generando las páginas de cada grado)
    for grado_num, data in grados_data.items():
        mat_data = data['mat']
        lec_data = data['lec']
        
        # PÁGINA A: PDF
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"INFORME DEL GRADO {grado_num} - ESTADÍSTICAS Y QUINTILES", fontsize=16, fontweight='bold', color='darkblue', y=0.96)
        
        part_text = "TASA DE FINALIZACIÓN DEL GRADO:\n\n"
        if lec_data: part_text += f"LENGUA: Registrados: {lec_data['part']['Registrados']:,}  |  100% Completos: {lec_data['part']['Completos']:,} ({lec_data['part']['Pct']:.1f}%)\n"
        if mat_data: part_text += f"MATEMÁTICA: Registrados: {mat_data['part']['Registrados']:,}  |  100% Completos: {mat_data['part']['Completos']:,} ({mat_data['part']['Pct']:.1f}%)\n"
        fig.text(0.5, 0.87, part_text, ha='center', va='center', fontsize=11, bbox=dict(facecolor='whitesmoke', edgecolor='gray', boxstyle='round,pad=0.5'))

        ax_desc = fig.add_axes([0.1, 0.60, 0.8, 0.15])
        desc_dfs = []
        if lec_data: desc_dfs.append(lec_data['desc'])
        if mat_data: desc_dfs.append(mat_data['desc'])
        if desc_dfs: format_table_pdf(ax_desc, pd.concat(desc_dfs, ignore_index=True), "Estadísticas Descriptivas (Alumnos 100%)")
        
        fig.text(0.5, 0.50, txt_quintiles, ha='center', va='top', fontsize=10, style='italic', color='black')

        ax_quint = fig.add_axes([0.1, 0.05, 0.8, 0.35])
        quint_dfs = []
        if lec_data: quint_dfs.append(lec_data['quint'])
        if mat_data: quint_dfs.append(mat_data['quint'])
        if quint_dfs: format_table_pdf(ax_quint, pd.concat(quint_dfs, ignore_index=True), "Distribución por Quintiles", bold_col_name="Rango")
        pdf.savefig(fig)
        plt.close()

        # PÁGINA: PDF - Histogramas Estudiantes (Estandarizado con Z-Score)
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"GRADO {grado_num} - DISTRIBUCIÓN DE ALUMNOS", fontsize=16, fontweight='bold', color='darkblue', y=0.95)
        fig.text(0.5, 0.88, txt_hist_alum, ha='center', va='top', fontsize=10, style='italic', color='black')
        gs_hist_al = fig.add_gridspec(1, 2, wspace=0.2, top=0.80, bottom=0.15)
        
        if lec_data:
            ax_hist_al_lec = fig.add_subplot(gs_hist_al[0])
            sns.histplot(lec_data['df_scores'], bins=20, color='#e34a33', edgecolor='black', ax=ax_hist_al_lec)
            ax_hist_al_lec.axvline(0, color='black', linestyle='--')
            ax_hist_al_lec.set_title("Z-score Estudiantes: Lengua")
        if mat_data:
            ax_hist_al_mat = fig.add_subplot(gs_hist_al[1] if lec_data else gs_hist_al[0])
            sns.histplot(mat_data['df_scores'], bins=20, color='#2c7fb8', edgecolor='black', ax=ax_hist_al_mat)
            ax_hist_al_mat.axvline(0, color='black', linestyle='--')
            ax_hist_al_mat.set_title("Z-score Estudiantes: Matemática")
        pdf.savefig(fig)
        plt.close()

        # PÁGINA: PDF - Histogramas Escuelas
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"GRADO {grado_num} - DISTRIBUCIÓN DE ESCUELAS", fontsize=16, fontweight='bold', color='darkblue', y=0.95)
        fig.text(0.5, 0.88, txt_hist, ha='center', va='top', fontsize=10, style='italic', color='black')
        gs_hist = fig.add_gridspec(1, 2, wspace=0.2, top=0.80, bottom=0.15)
        if lec_data:
            ax_hist_lec = fig.add_subplot(gs_hist[0])
            sns.histplot(lec_data['esc_z']['Z_Score'], bins=20, color='#e34a33', edgecolor='black', ax=ax_hist_lec)
            ax_hist_lec.axvline(0, color='black', linestyle='--')
            ax_hist_lec.set_title("Z-score Escuelas: Lengua")
        if mat_data:
            ax_hist_mat = fig.add_subplot(gs_hist[1] if lec_data else gs_hist[0])
            sns.histplot(mat_data['esc_z']['Z_Score'], bins=20, color='#2c7fb8', edgecolor='black', ax=ax_hist_mat)
            ax_hist_mat.axvline(0, color='black', linestyle='--')
            ax_hist_mat.set_title("Z-score Escuelas: Matemática")
        pdf.savefig(fig)
        plt.close()

        # PÁGINA C: PDF - Dificultad Ítems
        materias_validas = sum([1 for x in [lec_data, mat_data] if x is not None])
        if materias_validas > 0:
            fig, axes = plt.subplots(materias_validas, 1, figsize=(11, 8.5))
            fig.suptitle(f"GRADO {grado_num} - ANÁLISIS DE DIFICULTAD DE ÍTEMS", fontsize=16, fontweight='bold', color='darkblue', y=0.95)
            if materias_validas == 1: axes = [axes]
            
            idx = 0
            def plot_bars(ax, df, mat_name, color):
                sns.barplot(data=df, x='Item', y='Pct', color=color, ax=ax)
                ax.set_title(f"Porcentaje de Aciertos por Ítem: {mat_name}", fontweight='bold')
                ax.set_ylabel("Correctas (%)")
                ax.set_ylim(0, 115) 
                ax.tick_params(axis='x', rotation=90, labelsize=7)
                for container in ax.containers:
                    ax.bar_label(container, fmt='%.1f%%', padding=3, rotation=90, size=7, fontweight='bold')

            if lec_data: 
                plot_bars(axes[idx], lec_data['dif'], "Lengua", "#e34a33")
                idx += 1
            if mat_data: 
                plot_bars(axes[idx], mat_data['dif'], "Matemática", "#2c7fb8")
            plt.tight_layout(rect=[0, 0, 1, 0.9])
            pdf.savefig(fig)
            plt.close()

# =========================================================================
# 5. CREACIÓN DEL WORD EDITABLE
# =========================================================================
print("Generando Documento Editable en Word (DOCX)...\n")

def add_editable_table(doc, df, title):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Encabezados en negrita
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            
    # Filas (Resaltar "Rango" en negrita)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            if df.columns[i] == 'Rango':
                for run in row_cells[i].paragraphs[0].runs:
                    run.font.bold = True

def add_plot_to_word(doc, fig):
    mem_stream = io.BytesIO()
    fig.savefig(mem_stream, format='png', bbox_inches='tight')
    mem_stream.seek(0)
    doc.add_picture(mem_stream, width=Inches(6.0))
    mem_stream.close()

doc = Document()
doc.add_heading("RESULTADOS DESCRIPTIVOS DE LA APLICACIÓN\nPRUEBA DE RESULTADOS\nFEBRERO 2026", 0)
doc.add_page_break()

doc.add_heading("RESUMEN EJECUTIVO GLOBAL", level=1)
add_editable_table(doc, global_df, "Tasa de Participación Global (Todos los Grados)")
doc.add_paragraph(txt_boxplot.replace('\n', ' '))

# Generar solo el gráfico de boxplots para Word
if not df_global_estand.empty:
    materias_presentes = df_global_estand['Materia'].unique()
    fig, axes = plt.subplots(len(materias_presentes), 1, figsize=(8, 4 * len(materias_presentes)))
    if len(materias_presentes) == 1: axes = [axes]
    for i, mat in enumerate(materias_presentes):
        sns.boxplot(data=df_global_estand[df_global_estand['Materia'] == mat], x='Grado_Num', y='Puntaje_Transformado', color='lightgray', ax=axes[i])
        axes[i].set_title(f"Distribución {mat.upper()} (Z+5)")
        axes[i].set_ylim(0, 10)
    plt.tight_layout()
    add_plot_to_word(doc, fig)
    plt.close(fig)

doc.add_page_break()

for grado_num, data in grados_data.items():
    mat_data = data['mat']
    lec_data = data['lec']
    
    doc.add_heading(f"INFORME DEL GRADO {grado_num}", level=1)
    
    part_text = "TASA DE FINALIZACIÓN DEL GRADO:\n"
    if lec_data: part_text += f"LENGUA: Registrados: {lec_data['part']['Registrados']:,} | 100% Completos: {lec_data['part']['Completos']:,} ({lec_data['part']['Pct']:.1f}%)\n"
    if mat_data: part_text += f"MATEMÁTICA: Registrados: {mat_data['part']['Registrados']:,} | 100% Completos: {mat_data['part']['Completos']:,} ({mat_data['part']['Pct']:.1f}%)"
    doc.add_paragraph(part_text)

    # Tablas Editables Nativas
    desc_dfs = []
    if lec_data: desc_dfs.append(lec_data['desc'])
    if mat_data: desc_dfs.append(mat_data['desc'])
    if desc_dfs: add_editable_table(doc, pd.concat(desc_dfs, ignore_index=True), "Estadísticas Descriptivas (Alumnos 100%)")
    
    doc.add_paragraph(txt_quintiles.replace('\n', ' '))
    
    quint_dfs = []
    if lec_data: quint_dfs.append(lec_data['quint'])
    if mat_data: quint_dfs.append(mat_data['quint'])
    if quint_dfs: add_editable_table(doc, pd.concat(quint_dfs, ignore_index=True), "Distribución por Quintiles")

    # Gráficos de Histogramas de ALUMNOS para Word (Estandarizado)
    materias_validas = sum([1 for x in [lec_data, mat_data] if x is not None])
    doc.add_paragraph(txt_hist_alum.replace('\n', ' '))
    if materias_validas > 0:
        fig, axes = plt.subplots(1, materias_validas, figsize=(8, 4))
        if materias_validas == 1: axes = [axes]
        idx = 0
        if lec_data:
            sns.histplot(lec_data['df_scores'], bins=20, color='#e34a33', ax=axes[idx])
            axes[idx].axvline(0, color='black', linestyle='--')
            axes[idx].set_title("Z-score Estudiantes: Lengua")
            idx += 1
        if mat_data:
            sns.histplot(mat_data['df_scores'], bins=20, color='#2c7fb8', ax=axes[idx])
            axes[idx].axvline(0, color='black', linestyle='--')
            axes[idx].set_title("Z-score Estudiantes: Matemática")
        plt.tight_layout()
        add_plot_to_word(doc, fig)
        plt.close(fig)

    # Gráfico de Histogramas de ESCUELAS para Word
    doc.add_paragraph(txt_hist.replace('\n', ' '))
    if materias_validas > 0:
        fig, axes = plt.subplots(1, materias_validas, figsize=(8, 4))
        if materias_validas == 1: axes = [axes]
        idx = 0
        if lec_data:
            sns.histplot(lec_data['esc_z']['Z_Score'], bins=20, color='#e34a33', ax=axes[idx])
            axes[idx].axvline(0, color='black', linestyle='--')
            axes[idx].set_title("Z-score Escuelas: Lengua")
            idx += 1
        if mat_data:
            sns.histplot(mat_data['esc_z']['Z_Score'], bins=20, color='#2c7fb8', ax=axes[idx])
            axes[idx].axvline(0, color='black', linestyle='--')
            axes[idx].set_title("Z-score Escuelas: Matemática")
        plt.tight_layout()
        add_plot_to_word(doc, fig)
        plt.close(fig)

    # Gráfico de Barras para Word
    if materias_validas > 0:
        fig, axes = plt.subplots(materias_validas, 1, figsize=(8, 3.5 * materias_validas))
        if materias_validas == 1: axes = [axes]
        idx = 0
        def plot_bars_word(ax, df, mat_name, color):
            sns.barplot(data=df, x='Item', y='Pct', color=color, ax=ax)
            ax.set_title(f"Porcentaje de Aciertos: {mat_name}")
            ax.set_ylim(0, 115) 
            ax.tick_params(axis='x', rotation=90, labelsize=7)
            for container in ax.containers:
                ax.bar_label(container, fmt='%.1f%%', padding=3, rotation=90, size=7)

        if lec_data: 
            plot_bars_word(axes[idx], lec_data['dif'], "Lengua", "#e34a33")
            idx += 1
        if mat_data: 
            plot_bars_word(axes[idx], mat_data['dif'], "Matemática", "#2c7fb8")
        plt.tight_layout()
        add_plot_to_word(doc, fig)
        plt.close(fig)
        
    doc.add_page_break()

doc.save(archivo_docx)

# TXT OUTPUT
print_txt("=================================================================\n")
print_txt(" RESUMEN GLOBAL DE PARTICIPACIÓN (TODOS LOS GRADOS)\n")
print_txt("=================================================================\n")
for mat in ['Matemática', 'Lengua']:
    reg = global_totals[mat]['Registrados']
    com = global_totals[mat]['100%_Completos']
    pct = (com / reg) * 100 if reg > 0 else 0
    print_txt(f" {mat.upper()}:\n")
    print_txt(f"   - Gran Total Estudiantes Registrados : {reg:,}\n")
    print_txt(f"   - Gran Total con Prueba al 100%      : {com:,} ({pct:.1f}%)\n")
print_txt("=================================================================\n")

print("✅ PROCESO MAESTRO COMPLETADO EXITOSAMENTE.")