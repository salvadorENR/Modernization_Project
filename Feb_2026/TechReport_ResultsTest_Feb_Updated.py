import os
import datetime
import io
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.backends.backend_pdf import PdfPages
import seaborn as sns

# Eha'ã e-importa python-docx pe Word archívo ñemoheñóirã
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Instala la librería para Word usando el comando: pip install python-docx")
    exit()

# =========================================================================
# 1. CONFIGURACIÓN INICIAL Y CREACIÓN DE CARPETAS
# =========================================================================
script_dir = os.path.dirname(os.path.abspath(__file__))

dir_txt = os.path.join(script_dir, "Analisis_Diagnostico")
dir_pdf = os.path.join(script_dir, "Reporte_Tecnico")
dir_word = os.path.join(script_dir, "Reporte_Word")

os.makedirs(dir_txt, exist_ok=True)
os.makedirs(dir_pdf, exist_ok=True)
os.makedirs(dir_word, exist_ok=True)

archivo_pdf = os.path.join(dir_pdf, "Reporte_Tecnico_100_Completos.pdf")
archivo_txt_global = os.path.join(dir_txt, "Reporte_Completitud_100.txt")
archivo_word = os.path.join(dir_word, "Reporte_Resultados_Final.docx")

# --- DICCIONARIO DE ÍTEMS A OMITIR ---
ITEMS_A_ELIMINAR = {
    'Matemática': {
        3: ['MAT3992'],
        4: ['MAT3998', 'MAT4025', 'MAT4027'],
        5: ['MAT4039', 'MAT4043', 'MAT4044', 'MAT4050', 'MAT4055'],
        6: ['MAT4077'],
        7: ['MAT4112'],
        10: ['MAT4198', 'MAT4199', 'MAT4200', 'MAT4202', 'MAT4215'],
        11: ['MAT4228']
    },
    'Lengua': {
        3: ['LEC3154', 'LEC3157', 'LEC3168'],
        4: ['LEC3193', 'LEC3207'],
        6: ['LEC3262', 'LEC3267'],
        8: ['LEC3312'],
        10: ['LEC3366']
    }
}

global_totals = {
    'Matemática': {'Registrados': 0, '100%_Completos': 0},
    'Lengua': {'Registrados': 0, '100%_Completos': 0}
}
df_global_estand = pd.DataFrame()
tabla_resumen_items = {'Matemática': [], 'Lengua': []}
participacion_grados_mat = []
participacion_grados_lec = []

with open(archivo_txt_global, 'w', encoding='utf-8') as f:
    f.write("=================================================================\n")
    f.write(" REPORTE DE COMPLETITUD Y FILTRADO 100% (ANÁLISIS INDEPENDIENTE)\n")
    f.write("=================================================================\n\n")

def print_txt(texto):
    print(texto, end='')
    with open(archivo_txt_global, 'a', encoding='utf-8') as f:
        f.write(texto)

# =========================================================================
# 2. FUNCIONES DE PROCESAMIENTO Y FORMATO
# =========================================================================

def draw_silhouette(ax, x_offset, y_offset, face_color, edge_color, line_width):
    head = patches.Circle((x_offset, y_offset + 1.6), 0.2, linewidth=line_width, 
                          edgecolor=edge_color, facecolor=face_color, zorder=2)
    ax.add_patch(head)
    
    body_coords = [
        (x_offset - 0.2, y_offset + 0.5), 
        (x_offset + 0.2, y_offset + 0.5), 
        (x_offset + 0.25, y_offset + 1.4), 
        (x_offset - 0.25, y_offset + 1.4), 
    ]
    body = patches.Polygon(body_coords, closed=True, linewidth=line_width, 
                           edgecolor=edge_color, facecolor=face_color, 
                           joinstyle='round', zorder=1)
    ax.add_patch(body)

def generar_slide_pictogramas(pdf, materia, datos_tabla):
    if not datos_tabla: return
    fig, ax = plt.subplots(figsize=(11, 8.5))
    fig.suptitle(f"REPRESENTACIÓN GRÁFICA DE DESEMPEÑO: {materia.upper()}", fontsize=16, fontweight='bold', color='darkblue', y=0.95)
    fig.text(0.5, 0.90, "Proporción de estudiantes que alcanzan al menos el 60% de ítems válidos", ha='center', va='center', fontsize=12)

    y_max = (len(datos_tabla) - 1) * 1.8
    legend_y = y_max + 2.0
    
    draw_silhouette(ax, x_offset=3, y_offset=legend_y, face_color='#ADD8E6', edge_color='black', line_width=1)
    ax.text(3.5, legend_y + 1.15, "Alcanzó el 60% de ítems válidos", va='center', ha='left', fontsize=10)
    draw_silhouette(ax, x_offset=9, y_offset=legend_y, face_color='black', edge_color='black', line_width=1)
    ax.text(9.5, legend_y + 1.15, "No alcanzó el 60%", va='center', ha='left', fontsize=10)

    ax.set_xlim(0, 16)
    ax.set_ylim(0, legend_y + 2.5)
    ax.axis('off')

    for i, row in enumerate(datos_tabla):
        y_pos = (len(datos_tabla) - 1 - i) * 1.8
        ax.text(1.2, y_pos + 1.15, f"{row['Grado']}", va='center', ha='right', fontsize=11, fontweight='bold', color='#1e293b')
        
        aprobados = row['Aprobados_10']
        for j in range(10):
            x_off = 2.0 + j * 0.8
            if j < aprobados:
                draw_silhouette(ax, x_offset=x_off, y_offset=y_pos, face_color='#ADD8E6', edge_color='black', line_width=1)
            else:
                draw_silhouette(ax, x_offset=x_off, y_offset=y_pos, face_color='black', edge_color='black', line_width=1)
        
        frase_plot = f"{aprobados} de cada 10 estudiantes fueron\ncapaces de responder correctamente\nal menos el 60% de los ítems válidos."
        ax.text(10.2, y_pos + 1.15, frase_plot, va='center', ha='left', fontsize=9, color='#002d5a', style='italic')
        
        if i < len(datos_tabla) - 1:
            ax.hlines(y=y_pos + 0.1, xmin=0.5, xmax=15.5, color='#cbd5e1', linewidth=1, linestyle='--')

    pdf.savefig(fig)
    plt.close()

def format_table_pdf(ax, df, title=""):
    ax.axis('tight')
    ax.axis('off')
    if title:
        ax.set_title(title, fontweight='bold', pad=12, fontsize=12, color='darkblue')
    table = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.7) 
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor('darkblue')
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')
        else:
            if col >= 0 and df.columns[col] == 'Rango':
                cell.get_text().set_fontweight('bold')

def format_table_pdf_part(ax, df, title=""):
    ax.axis('tight')
    ax.axis('off')
    if title:
        # Pad estándar. La cercanía al título se controla con la caja de add_axes ahora.
        ax.set_title(title, fontweight='bold', pad=10, fontsize=12, color='darkblue')
    
    # loc='center' es mucho más seguro para la visualización estable de la tabla.
    table = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.7) 
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor('darkblue')
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')

def format_table_pdf_items(ax, df, title=""):
    ax.axis('tight')
    ax.axis('off')
    if title:
        ax.set_title(title, fontweight='bold', pad=15, fontsize=14, color='darkblue')
    table = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.8) 
    
    col_widths = [0.1, 0.2, 0.2, 0.5] 
    for (row, col), cell in table.get_celld().items():
        cell.set_width(col_widths[col])
        if row == 0:
            cell.set_facecolor('darkblue')
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')
        else:
            if col == 3: 
                cell.get_text().set_color('#b91c1c') 
                cell.get_text().set_fontsize(9)

def agregar_tabla_word(doc, df, titulo=""):
    if titulo:
        doc.add_heading(titulo, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            if df.columns[i] == 'Rango':
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    doc.add_paragraph("\n")

def guardar_grafico_word(doc, fig, ancho=6):
    mem_buf = io.BytesIO()
    fig.savefig(mem_buf, format='png', bbox_inches='tight')
    doc.add_picture(mem_buf, width=Inches(ancho))
    mem_buf.close()

def leer_y_limpiar(filepath):
    if not os.path.exists(filepath): return None
    try:
        df = pd.read_csv(filepath, sep='|', encoding='utf-8', dtype=str, on_bad_lines='skip')
    except UnicodeDecodeError:
        df = pd.read_csv(filepath, sep='|', encoding='latin-1', dtype=str, on_bad_lines='skip')
    df.columns = [c.replace('"', '').strip() for c in df.columns]
    if 'Documento' not in df.columns: return None
    df['Documento'] = df['Documento'].str.strip()
    if 'Centro' in df.columns:
        df['Centro'] = df['Centro'].str.strip().replace('', 'DESCONOCIDO').fillna('DESCONOCIDO')
    else:
        df['Centro'] = 'DESCONOCIDO'
    df.dropna(subset=['Documento'], inplace=True)
    df = df[df['Documento'] != '']
    df.drop_duplicates(subset=['Documento'], keep='first', inplace=True)
    return df

def procesar_100_porciento(df, materia, prefijo, grado_num, items_eliminados_list):
    if df is None or df.empty: return None
    
    item_cols = [c for c in df.columns if c.startswith(prefijo) and c[len(prefijo):].isdigit()]
    tot_items_validos = len(item_cols)
    if tot_items_validos == 0: return None
    
    valid_mask = df[item_cols].isin(['0', '1']).sum(axis=1) == tot_items_validos
    df_100 = df[valid_mask].copy()
    
    df_100['Puntaje_Total'] = df_100[item_cols].astype(int).sum(axis=1)
    
    tot_items_originales = tot_items_validos + len(items_eliminados_list)
    umbral_60 = tot_items_validos * 0.6
    pct_60 = (df_100['Puntaje_Total'] >= umbral_60).mean() * 100
    
    x_de_10 = int(round(pct_60 / 10))
    if x_de_10 == 0 and pct_60 > 0: x_de_10 = 1 
    
    excluidos_str = ", ".join(items_eliminados_list) if items_eliminados_list else "Ninguno"
    
    tabla_resumen_items[materia].append({
        'Grado': f"Grado {grado_num}°",
        'Ítems Originales': tot_items_originales,
        'Ítems Válidos': tot_items_validos,
        'Ítems Excluidos': excluidos_str,
        'Aprobados_10': x_de_10 
    })
    
    tot_inicial, tot_final = len(df), len(df_100)
    pct = (tot_final / tot_inicial) * 100 if tot_inicial > 0 else 0
    global_totals[materia]['Registrados'] += tot_inicial
    global_totals[materia]['100%_Completos'] += tot_final
    
    print_txt(f" [{materia.upper()}]\n")
    print_txt(f"   - Total de ítems válidos        : {tot_items_validos}\n")
    print_txt(f"   - Estudiantes Registrados       : {tot_inicial:,}\n")
    print_txt(f"   - Completaron el 100%           : {tot_final:,} ({pct:.1f}%)\n\n")
    
    participacion = {'Registrados': tot_inicial, 'Completos': tot_final, 'Pct': pct}
    return {'df': df_100, 'items': item_cols, 'participacion': participacion}

def extraer_metricas(datos_obj, materia, grado_num):
    global df_global_estand
    if datos_obj is None or datos_obj['df'].empty: return None
    
    df_100, item_cols, participacion = datos_obj['df'], datos_obj['items'], datos_obj['participacion']
    
    media = df_100['Puntaje_Total'].mean()
    sd = df_100['Puntaje_Total'].std(ddof=1)
    if pd.isna(sd) or sd == 0: sd = 1
    
    df_100['Z_Score'] = (df_100['Puntaje_Total'] - media) / sd
    df_100['Puntaje_Transformado'] = df_100['Z_Score'] + 5
    
    df_temp = df_100[['Centro', 'Documento', 'Puntaje_Transformado']].copy()
    df_temp['Materia'] = materia
    df_temp['Grado_Num'] = grado_num
    df_global_estand = pd.concat([df_global_estand, df_temp], ignore_index=True)
    
    desc = pd.DataFrame([{
        'Materia': materia, 'N (100%)': f"{len(df_100):,}", 'Media': round(media, 2),
        'Mediana': df_100['Puntaje_Total'].median(), 'DS': round(sd, 2),
        'Mínimo': df_100['Puntaje_Total'].min(), 'Máximo': df_100['Puntaje_Total'].max()
    }])
    
    df_100['Quintil'] = pd.qcut(df_100['Puntaje_Total'].rank(method='first'), 5, labels=[1, 2, 3, 4, 5])
    quint_list = []
    for q, g in df_100.groupby('Quintil', observed=True):
        quint_list.append({
            'Materia': materia, 'Quintil': q,
            'Rango': f"{g['Puntaje_Total'].min():.1f} - {g['Puntaje_Total'].max():.1f}",
            'N': f"{len(g):,}", '%': f"{(len(g) / len(df_100) * 100):.1f}%"
        })
    quint = pd.DataFrame(quint_list)
    
    esc_z = df_100.groupby('Centro')['Puntaje_Total'].mean().reset_index()
    esc_mean = esc_z['Puntaje_Total'].mean()
    esc_sd = esc_z['Puntaje_Total'].std(ddof=1)
    if pd.isna(esc_sd) or esc_sd == 0: esc_sd = 1
    esc_z['Z_Score'] = (esc_z['Puntaje_Total'] - esc_mean) / esc_sd
    
    dif_series = (df_100[item_cols] == '1').mean() * 100
    dif = pd.DataFrame({'Item': dif_series.index, 'Pct': dif_series.values})
    
    return {'desc': desc, 'quint': quint, 'esc_z': esc_z, 'dif': dif, 'part': participacion, 'df': df_100}

# =========================================================================
# 3. FASE DE PROCESAMIENTO DE DATOS
# =========================================================================
print("\nIniciando Procesamiento de Datos...\n")
grados_data = {}

for grado_num in range(3, 12):
    archivo_mat = os.path.join(script_dir, f"Mat_{grado_num}.txt")
    archivo_lec = os.path.join(script_dir, f"Lec_{grado_num}.txt")
    
    df_mat_raw = leer_y_limpiar(archivo_mat)
    df_lec_raw = leer_y_limpiar(archivo_lec)
    if df_mat_raw is None and df_lec_raw is None: continue
        
    print_txt(f"-> Procesando Grado {grado_num}...\n")
    
    items_elim_mat = ITEMS_A_ELIMINAR['Matemática'].get(grado_num, [])
    items_elim_lec = ITEMS_A_ELIMINAR['Lengua'].get(grado_num, [])
    
    if df_mat_raw is not None and items_elim_mat:
        cols_to_drop = [c for c in items_elim_mat if c in df_mat_raw.columns]
        if cols_to_drop:
            df_mat_raw.drop(columns=cols_to_drop, inplace=True)

    if df_lec_raw is not None and items_elim_lec:
        cols_to_drop = [c for c in items_elim_lec if c in df_lec_raw.columns]
        if cols_to_drop:
            df_lec_raw.drop(columns=cols_to_drop, inplace=True)

    datos_mat = procesar_100_porciento(df_mat_raw, "Matemática", "MAT", grado_num, items_elim_mat)
    datos_lec = procesar_100_porciento(df_lec_raw, "Lengua", "LEC", grado_num, items_elim_lec)
    
    # Recolectar datos para tablas de participación por grado
    if datos_mat:
        participacion_grados_mat.append({
            'Grado': f"{grado_num}°",
            'Total Registrados': f"{datos_mat['participacion']['Registrados']:,}",
            '100% Completos': f"{datos_mat['participacion']['Completos']:,}",
            'Porcentaje': f"{datos_mat['participacion']['Pct']:.1f}%"
        })
    if datos_lec:
        participacion_grados_lec.append({
            'Grado': f"{grado_num}°",
            'Total Registrados': f"{datos_lec['participacion']['Registrados']:,}",
            '100% Completos': f"{datos_lec['participacion']['Completos']:,}",
            'Porcentaje': f"{datos_lec['participacion']['Pct']:.1f}%"
        })

    met_mat = extraer_metricas(datos_mat, "Matemática", grado_num)
    met_lec = extraer_metricas(datos_lec, "Lengua", grado_num)
    
    if met_mat or met_lec:
        grados_data[grado_num] = {'mat': met_mat, 'lec': met_lec}

# =========================================================================
# 4. CREACIÓN DEL PDF Y WORD CON COMENTARIOS INTEGRADOS
# =========================================================================
print("\nGenerando Documentos (PDF y Word)...\n")

with PdfPages(archivo_pdf) as pdf:
    # --- PÁGINA 1: PORTADA PDF ---
    fig = plt.figure(figsize=(11, 8.5))
    plt.axis('off')
    plt.text(0.5, 0.65, "RESULTADOS DESCRIPTIVOS\nDE LA APLICACIÓN", ha='center', va='center', fontsize=28, fontweight='bold', color='darkblue')
    plt.text(0.5, 0.45, "PRUEBA DE RESULTADOS", ha='center', va='center', fontsize=18)
    plt.text(0.5, 0.35, "FEBRERO 2026", ha='center', va='center', fontsize=14, style='italic', color='gray')
    pdf.savefig(fig)
    plt.close()

    # --- PÁGINA 2: RESUMEN DE PARTICIPACIÓN ---
    fig = plt.figure(figsize=(11, 8.5))
    fig.suptitle("RESUMEN GLOBAL DE PARTICIPACIÓN", fontsize=22, fontweight='bold', color='darkblue', y=0.94)
    
    # Tabla Global (Arriba)
    ax_global = fig.add_axes([0.15, 0.74, 0.7, 0.12])
    global_df = pd.DataFrame([
        {'Materia': 'Matemática', 
         'Total Registrados': f"{global_totals['Matemática']['Registrados']:,}", 
         'Completaron 100%': f"{global_totals['Matemática']['100%_Completos']:,}", 
         'Porcentaje': f"{(global_totals['Matemática']['100%_Completos']/max(1, global_totals['Matemática']['Registrados']))*100:.1f}%"},
        {'Materia': 'Lengua', 
         'Total Registrados': f"{global_totals['Lengua']['Registrados']:,}", 
         'Completaron 100%': f"{global_totals['Lengua']['100%_Completos']:,}", 
         'Porcentaje': f"{(global_totals['Lengua']['100%_Completos']/max(1, global_totals['Lengua']['Registrados']))*100:.1f}%"}
    ])
    format_table_pdf(ax_global, global_df, "Tasa de Participación Global (Todos los Grados)")

    # Tabla Matemática por Grados (Izquierda). Cajas limitadas a 0.48 de altura para que los títulos queden perfectos.
    if participacion_grados_mat:
        ax_mat_part = fig.add_axes([0.05, 0.15, 0.42, 0.48])
        df_mat_part = pd.DataFrame(participacion_grados_mat)
        format_table_pdf_part(ax_mat_part, df_mat_part, "Matemática: Participación por Grado")

    # Tabla Lengua por Grados (Derecha)
    if participacion_grados_lec:
        ax_lec_part = fig.add_axes([0.53, 0.15, 0.42, 0.48])
        df_lec_part = pd.DataFrame(participacion_grados_lec)
        format_table_pdf_part(ax_lec_part, df_lec_part, "Lengua: Participación por Grado")

    pdf.savefig(fig)
    plt.close()

    # --- PÁGINA 3: BOXPLOTS GLOBALES ---
    fig = plt.figure(figsize=(11, 8.5))
    fig.suptitle("RESUMEN GLOBAL DE DISTRIBUCIÓN", fontsize=22, fontweight='bold', color='darkblue', y=0.95)
    
    txt_global = ("El siguiente gráfico presenta la distribución de los puntajes estandarizados (escala con media 5\n"
                  "y desviación estándar 1, agrupando la mayor parte de los datos entre 2 y 8) para todos los grados evaluados.\n"
                  "Las cajas muestran el rango intercuartílico, permitiendo observar la variabilidad y la\n"
                  "tendencia central de la muestra evaluada.")
    fig.text(0.5, 0.88, txt_global, ha='center', va='center', fontsize=10)
    
    if not df_global_estand.empty:
        materias_presentes = df_global_estand['Materia'].unique()
        num_mats = len(materias_presentes)
        axes = [fig.add_axes([0.1, 0.50, 0.8, 0.28]), fig.add_axes([0.1, 0.08, 0.8, 0.28])] if num_mats == 2 else [fig.add_axes([0.1, 0.15, 0.8, 0.55])]
        
        for i, mat in enumerate(materias_presentes):
            ax = axes[i]
            df_plot = df_global_estand[df_global_estand['Materia'] == mat]
            sns.boxplot(data=df_plot, x='Grado_Num', y='Puntaje_Transformado', color='gray', ax=ax)
            ax.set_title(f"Distribución {mat.upper()} (Z+5)", fontweight='bold', pad=10)
            
            # Asegurar que el label 'Grado Analizado' se mantenga visualmente en ambos gráficos
            ax.set_xlabel("Grado Analizado")
            ax.set_ylabel("Puntaje Estandarizado")
            ax.set_ylim(0, 10)
            ax.grid(axis='y', linestyle='--', alpha=0.7)
    pdf.savefig(fig)
    plt.close()

    # --- PÁGINA 4: RESUMEN DE ÍTEMS EVALUADOS Y EXCLUIDOS (JUNTOS) ---
    if tabla_resumen_items['Matemática'] or tabla_resumen_items['Lengua']:
        fig = plt.figure(figsize=(11, 8.5))
        fig.text(0.5, 0.94, "RESUMEN DE ÍTEMS EVALUADOS Y EXCLUIDOS POR GRADO", ha='center', va='center', fontsize=20, fontweight='bold', color='darkblue')
        
        if tabla_resumen_items['Matemática']:
            ax_mat_tbl = fig.add_axes([0.05, 0.53, 0.9, 0.35])
            df_tbl_mat = pd.DataFrame(tabla_resumen_items['Matemática']).drop(columns=['Aprobados_10'])
            format_table_pdf_items(ax_mat_tbl, df_tbl_mat, "Área: Matemática")
            
        if tabla_resumen_items['Lengua']:
            ax_lec_tbl = fig.add_axes([0.05, 0.08, 0.9, 0.35])
            df_tbl_lec = pd.DataFrame(tabla_resumen_items['Lengua']).drop(columns=['Aprobados_10'])
            format_table_pdf_items(ax_lec_tbl, df_tbl_lec, "Área: Lengua")
            
        pdf.savefig(fig)
        plt.close()
        
    # --- PÁGINAS 5A/5B: PICTOGRAMAS ---
    if tabla_resumen_items['Matemática']:
        generar_slide_pictogramas(pdf, "Matemática", tabla_resumen_items['Matemática'])
    if tabla_resumen_items['Lengua']:
        generar_slide_pictogramas(pdf, "Lengua", tabla_resumen_items['Lengua'])

    # --- DETALLE POR GRADO ---
    for grado_num, data in grados_data.items():
        met_mat = data['mat']
        met_lec = data['lec']

        # PÁGINA A: Estadísticas y Quintiles
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"INFORME DEL GRADO {grado_num} - ESTADÍSTICAS Y QUINTILES", fontsize=16, fontweight='bold', color='darkblue', y=0.96)
        
        part_text = "TASA DE FINALIZACIÓN DEL GRADO:\n"
        if met_lec: part_text += f"LENGUA: Registrados: {met_lec['part']['Registrados']:,}  |  100% Completos: {met_lec['part']['Completos']:,} ({met_lec['part']['Pct']:.1f}%)\n"
        if met_mat: part_text += f"MATEMÁTICA: Registrados: {met_mat['part']['Registrados']:,}  |  100% Completos: {met_mat['part']['Completos']:,} ({met_mat['part']['Pct']:.1f}%)"
        fig.text(0.5, 0.88, part_text, ha='center', va='center', fontsize=10, bbox=dict(facecolor='whitesmoke', edgecolor='gray', boxstyle='round,pad=0.5'))

        ax_desc = fig.add_axes([0.1, 0.65, 0.8, 0.15]) 
        desc_dfs = []
        if met_lec: desc_dfs.append(met_lec['desc'])
        if met_mat: desc_dfs.append(met_mat['desc'])
        if desc_dfs: format_table_pdf(ax_desc, pd.concat(desc_dfs, ignore_index=True), "Estadísticas Descriptivas (Alumnos 100%)")

        txt_quint = ("La distribución por quintiles divide a la población estudiantil en cinco grupos de igual tamaño\n"
                     "(20% cada uno) ordenados según su puntaje.\nLos rangos muestran el nivel de desempeño alcanzado en cada estrato.")
        fig.text(0.5, 0.54, txt_quint, ha='center', va='center', fontsize=10) 

        ax_quint = fig.add_axes([0.1, 0.08, 0.8, 0.38]) 
        quint_dfs = []
        if met_lec: quint_dfs.append(met_lec['quint'])
        if met_mat: quint_dfs.append(met_mat['quint'])
        if quint_dfs: format_table_pdf(ax_quint, pd.concat(quint_dfs, ignore_index=True), "Distribución por Quintiles")
        pdf.savefig(fig)
        plt.close()

        # PÁGINA B: Histograma de ESTUDIANTES
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"GRADO {grado_num} - DISTRIBUCIÓN DE ALUMNOS", fontsize=16, fontweight='bold', color='darkblue', y=0.96)
        
        txt_est = ("Los histogramas muestran la distribución estandarizada (Z-score) de los puntajes obtenidos por todos los\n"
                   "estudiantes del grado.\nUn valor de 0 indica que el alumno se encuentra exactamente en el promedio general.")
        fig.text(0.5, 0.88, txt_est, ha='center', va='center', fontsize=10)

        gs_hist_st = fig.add_gridspec(1, 2, wspace=0.2, top=0.78, bottom=0.10)
        
        if met_lec:
            ax_lec_st = fig.add_subplot(gs_hist_st[0])
            sns.histplot(met_lec['df']['Z_Score'], bins=20, color='#e34a33', edgecolor='black', ax=ax_lec_st)
            mean_lec_st = met_lec['df']['Z_Score'].mean()
            ax_lec_st.axvline(mean_lec_st, color='black', linestyle=':', linewidth=2)
            ax_lec_st.set_title(f"Z-score Estudiantes: Lengua")
            ax_lec_st.set_xlabel("Z_Score")
            ax_lec_st.set_ylabel("Count")

        if met_mat:
            ax_mat_st = fig.add_subplot(gs_hist_st[1] if met_lec else gs_hist_st[0])
            sns.histplot(met_mat['df']['Z_Score'], bins=20, color='#2c7fb8', edgecolor='black', ax=ax_mat_st)
            mean_mat_st = met_mat['df']['Z_Score'].mean()
            ax_mat_st.axvline(mean_mat_st, color='black', linestyle=':', linewidth=2)
            ax_mat_st.set_title(f"Z-score Estudiantes: Matemática")
            ax_mat_st.set_xlabel("Z_Score")
            ax_mat_st.set_ylabel("Count")
            
        pdf.savefig(fig)
        plt.close()

        # PÁGINA C: Histograma de ESCUELAS
        fig = plt.figure(figsize=(11, 8.5))
        fig.suptitle(f"GRADO {grado_num} - DISTRIBUCIÓN DE ESCUELAS", fontsize=16, fontweight='bold', color='darkblue', y=0.96)
        
        txt_esc = ("Los histogramas muestran la distribución de los promedios de cada Centro Educativo en puntajes Z.\n"
                   "Un valor de 0 indica que la escuela se encuentra exactamente en el promedio global del grado.\n"
                   "Las barras ilustran la concentración de las escuelas evaluadas.")
        fig.text(0.5, 0.87, txt_esc, ha='center', va='center', fontsize=10)

        gs_hist = fig.add_gridspec(1, 2, wspace=0.2, top=0.77, bottom=0.10)
        
        if met_lec:
            ax_hist_lec = fig.add_subplot(gs_hist[0])
            sns.histplot(met_lec['esc_z']['Z_Score'], bins=20, color='#e34a33', edgecolor='black', ax=ax_hist_lec)
            mean_lec_esc = met_lec['esc_z']['Z_Score'].mean()
            ax_hist_lec.axvline(mean_lec_esc, color='black', linestyle=':', linewidth=2)
            ax_hist_lec.set_title(f"Z-score Escuelas: Lengua")
            ax_hist_lec.set_xlabel("Z_Score")
            ax_hist_lec.set_ylabel("Count")

        if met_mat:
            ax_hist_mat = fig.add_subplot(gs_hist[1] if met_lec else gs_hist[0])
            sns.histplot(met_mat['esc_z']['Z_Score'], bins=20, color='#2c7fb8', edgecolor='black', ax=ax_hist_mat)
            mean_mat_esc = met_mat['esc_z']['Z_Score'].mean()
            ax_hist_mat.axvline(mean_mat_esc, color='black', linestyle=':', linewidth=2)
            ax_hist_mat.set_title(f"Z-score Escuelas: Matemática")
            ax_hist_mat.set_xlabel("Z_Score")
            ax_hist_mat.set_ylabel("Count")

        pdf.savefig(fig)
        plt.close()

        # PÁGINA D: Barras de Dificultad
        materias_validas = sum([1 for x in [met_lec, met_mat] if x is not None])
        if materias_validas > 0:
            fig, axes = plt.subplots(materias_validas, 1, figsize=(11, 8.5))
            
            fig.suptitle(f"GRADO {grado_num} - PORCENTAJE DE RESPUESTAS CORRECTAS POR ÍTEM", fontsize=16, fontweight='bold', color='darkblue', y=0.95)
            
            if materias_validas == 1: axes = [axes]
            
            idx = 0
            def plot_bars(ax, df, mat_name, color):
                sns.barplot(data=df, x='Item', y='Pct', color=color, ax=ax)
                ax.set_title(f"Porcentaje de Aciertos por Ítem: {mat_name}", fontweight='bold')
                ax.set_ylabel("Correctas (%)")
                ax.set_xlabel("Item")
                ax.set_ylim(0, 115) 
                ax.set_yticks(np.arange(0, 101, 20))
                ax.tick_params(axis='x', rotation=90, labelsize=7)
                for container in ax.containers:
                    ax.bar_label(container, fmt='%.1f%%', padding=3, rotation=90, size=7, fontweight='bold')

            if met_lec: 
                plot_bars(axes[idx], met_lec['dif'], "Lengua", "#e34a33")
                idx += 1
            if met_mat: 
                plot_bars(axes[idx], met_mat['dif'], "Matemática", "#2c7fb8")

            plt.tight_layout(rect=[0, 0, 1, 0.9])
            pdf.savefig(fig)
            plt.close()

# =========================================================================
# 5. WORD REPORTE ÑEMOHEÑÓI
# =========================================================================
doc = Document()
doc.add_heading('RESULTADOS DESCRIPTIVOS DE LA APLICACIÓN', 0)
doc.add_paragraph('PRUEBA DE RESULTADOS')
doc.add_paragraph('FEBRERO 2026')

doc.add_heading('1. Resumen Global de Participación', level=1)
# Tabla Global Word
agregar_tabla_word(doc, global_df, "Tasa de Participación Global")
# Tablas Grados Word
if participacion_grados_mat:
    agregar_tabla_word(doc, pd.DataFrame(participacion_grados_mat), "Matemática: Participación por Grado")
if participacion_grados_lec:
    agregar_tabla_word(doc, pd.DataFrame(participacion_grados_lec), "Lengua: Participación por Grado")

doc.add_heading('2. Distribución Global de Puntajes', level=1)
if not df_global_estand.empty:
    fig_nac, ax_nac = plt.subplots(figsize=(8, 4))
    sns.boxplot(data=df_global_estand, x='Grado_Num', y='Puntaje_Transformado', hue='Materia', palette=['gray', 'gray'], ax=ax_nac)
    ax_nac.set_title("Distribución Nacional (Puntaje Z+5)")
    guardar_grafico_word(doc, fig_nac)
    plt.close(fig_nac)

doc.add_heading('3. Resumen de Ítems Evaluados y Excluidos', level=1)
if tabla_resumen_items['Matemática']:
    df_word_mat = pd.DataFrame(tabla_resumen_items['Matemática']).drop(columns=['Aprobados_10'])
    agregar_tabla_word(doc, df_word_mat, "Área: Matemática")
if tabla_resumen_items['Lengua']:
    df_word_lec = pd.DataFrame(tabla_resumen_items['Lengua']).drop(columns=['Aprobados_10'])
    agregar_tabla_word(doc, df_word_lec, "Área: Lengua")

doc.add_heading('4. Análisis Detallado por Grado', level=1)
for g, data in grados_data.items():
    doc.add_page_break()
    doc.add_heading(f'GRADO {g}', level=2)
    for m_key in ['mat', 'lec']:
        if data[m_key]:
            m_name = "Matemática" if m_key == 'mat' else "Lengua"
            doc.add_heading(f'Sub-área: {m_name}', level=3)
            agregar_tabla_word(doc, data[m_key]['quint'], f"Distribución de Quintiles {m_name}")
            
            fig_dif, ax_dif = plt.subplots(figsize=(10, 3.5))
            sns.barplot(data=data[m_key]['dif'], x='Item', y='Pct', color='skyblue' if m_key=='mat' else 'salmon', ax=ax_dif)
            ax_dif.tick_params(axis='x', rotation=90, labelsize=7)
            ax_dif.set_title(f"Porcentaje de Respuestas Correctas por Ítem - {m_name}")
            guardar_grafico_word(doc, fig_dif, ancho=6.5)
            plt.close(fig_dif)

doc.save(archivo_word)

# =========================================================================
# 6. TXT LOG ÑEMOPAHA
# =========================================================================
print_txt("=================================================================\n")
print_txt(" RESUMEN GLOBAL DE PARTICIPACIÓN (TODOS LOS GRADOS)\n")
print_txt("=================================================================\n")
for mat in ['Matemática', 'Lengua']:
    reg = global_totals[mat]['Registrados']
    com = global_totals[mat]['100%_Completos']
    pct = (com / reg) * 100 if reg > 0 else 0
    print_txt(f" {mat.upper()}:\n")
    print_txt(f"   - Gran Total Estudiantes Registrados : {reg:,}\n")
    print_txt(f"   - Gran Total con Prueba al 100%      : {com:,} ({pct:.1f}%)\n")
print_txt("=================================================================\n")

print(f"✅ Documentos generados correctamente en sus carpetas respectivas.")